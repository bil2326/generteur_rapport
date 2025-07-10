import pickle
from docx import Document
from docx.shared import Inches
from datetime import datetime
from api_calls import build_content
from datetime import date
import io

class Report:
    def __init__(self, report_id, report_subject, intervention_date, intervention_zone, site):
        self.report_id = f"DE / ETUDE-01809 / {report_id}"
        self.report_subject = report_subject
        self.intervention_date = intervention_date
        self.intervention_zone = intervention_zone
        self.site = site
        self.content = {
            "etat_des_lieux": [],
            "diagnostic": [],
            "preconisation": []
        }
        self.images = []
        self.vocals = []


    def add_image(self, image_bytes):
        self.images.append(image_bytes)

    def add_vocal(self, vocal_path):
        self.vocals.append(vocal_path)

    def write_content(self, image_bytes, vocal_path):

        section_result = build_content(image_bytes, vocal_path)

        # parse result json string if needed
        self.content["etat_des_lieux"].append(section_result["etat_des_lieux"])
        self.content["diagnostic"].append(section_result["diagnostic"])
        self.content["preconisation"].append(section_result["preconisation"])

    def save_as_pkl(self, path):
        with open(path, "wb") as f:
            pickle.dump(self, f)

    @staticmethod
    def read_pkl_file(path):
        with open(path, "rb") as f:
            return pickle.load(f)

    def export_report_as_docx_file(self, path):

        doc = Document("template.docx")
        remplacement = {}
        remplacement[r"{{report_subject}}"] = self.report_subject
        remplacement[r"{{intervention_date}}"] = str(self.intervention_date)
        remplacement[r"{{intervention_zone}}"] = self.intervention_zone
        remplacement[r"{{site}}"] = self.site

        remplacement[r"{{date_now}}"] = str(date.today())
        remplacement[r"{{etat_des_lieux}}"] = '+ ' + "\n+ ".join(self.content["etat_des_lieux"])
        remplacement[r"{{diagnostic}}"] = '+ ' + "\n+ ".join(self.content["diagnostic"])
        remplacement[r"{{preconisation}}"] = '+ ' + "\n+ ".join(self.content["preconisation"])
        doc = self.replace_placeholder(doc, remplacement)

        doc.save(path)





    def replace_placeholder(self, doc, remplacements):
        
        if self.images:
            number_of_images = len(self.images)
            for index_image in range(0, number_of_images):
                image_tag = f"{{image_{index_image}}}"
                img_bytes = self.images[index_image]

                for paragraph in doc.paragraphs:
                    if image_tag in paragraph.text:
                        paragraph.text = paragraph.text.replace(image_tag, "")
                        
                        run = paragraph.add_run()
  
                        image_stream = io.BytesIO(img_bytes)
                        run.add_picture(image_stream, width=Inches(4))
                        if index_image < len(self.images) - 1:
                            paragraph = doc.add_paragraph()
                        break


        for paragraphe in doc.paragraphs:

            for marqueur, valeur in remplacements.items():

                if marqueur in paragraphe.text:
                    paragraphe.text = paragraphe.text.replace(marqueur, valeur)


        for table in doc.tables:
            for row in table.rows:
                try:
                    # Certaines lignes peuvent lever une exception ici à cause de cellules fusionnées verticalement
                    for cell in row.cells:
                        for marqueur, valeur in remplacements.items():
                            if marqueur in cell.text:
                                cell.text = cell.text.replace(marqueur, valeur)
                except ValueError:
                    # Ligne avec cellules fusionnées verticalement non directement accessibles
                    continue

        return doc




    def __str__(self):
        return f"{self.content['etat_des_lieux']}"