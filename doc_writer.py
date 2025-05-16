from docx import Document
from io import BytesIO
# from docx_custom_properties import CustomProperties


def replace_placeholder(doc, remplacements):

    for paragraphe in doc.paragraphs:

        for marqueur, valeur in remplacements.items():

            if marqueur in paragraphe.text:
                paragraphe.text = paragraphe.text.replace(marqueur, valeur)


    for table in doc.tables:
        for row in table.rows:
            try:
                # Certaines lignes peuvent lever une exception ici à cause de cellules fusionnées verticalement
                for cell in row.cells:
                    for marqueur, valeur in remplacements.items():
                        if marqueur in cell.text:
                            cell.text = cell.text.replace(marqueur, valeur)
            except ValueError:
                # Ligne avec cellules fusionnées verticalement non directement accessibles
                continue
    # props = CustomProperties(doc)

    # Injecte les valeurs dans les QuickParts
    doc.custom_properties["INTITULE DE L'AFFAIRE"] = remplacements["{{SITE}}"]
    doc.custom_properties["INTITULE DU DOCUMENT"] = f"{remplacements['{{zone}}']} – {remplacements['{{sujet}}']}"
    return doc


def init_doc():
    return Document('template.docx')



# def generateur_introduction(doc, sujet, date_intervention, zone, site):
#     doc.add_heading("Rapport d'Intervention", 0)
#     doc.add_heading("I) Introduction", level=1)
#     doc.add_paragraph(f"📌 Sujet : {sujet}")
#     doc.add_paragraph(f"📅 Date : {date_intervention.strftime('%d/%m/%Y')}")
#     doc.add_paragraph(f"🌍 Zone d’intervention : {zone}")
#     doc.add_paragraph(f"🏗️ Site : {site}")

#     return doc


def get_buffer(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

