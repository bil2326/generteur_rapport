from docx import Document
from io import BytesIO

def init_doc():
    return Document()

def generateur_introduction(doc, sujet, date_intervention, zone, site):
    doc.add_heading("Rapport d'Intervention", 0)
    doc.add_heading("I) Introduction", level=1)
    doc.add_paragraph(f"📌 Sujet : {sujet}")
    doc.add_paragraph(f"📅 Date : {date_intervention.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"🌍 Zone d’intervention : {zone}")
    doc.add_paragraph(f"🏗️ Site : {site}")

    return doc


def get_buffer(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer