import streamlit as st
from docx import Document
from io import BytesIO
import datetime

st.set_page_config(page_title="Générateur de Rapport", layout="centered")

st.subheader("🧩 I) Introduction")

# Champs du formulaire
sujet = st.text_input("📌 Sujet du rapport")

date_intervention = st.date_input("📅 Date d'intervention", value=datetime.date.today())

zone = st.text_area("🌍 Zone d’intervention")

site = st.text_area("🏗️ Site")

# Bouton pour générer
if st.button("📄 Générer le rapport"):
    if not sujet or not zone or not site:
        st.warning("Veuillez remplir tous les champs avant de générer le rapport.")
    else:
        # Création du document Word
        doc = Document()
        doc.add_heading("Rapport d'Intervention", 0)

        doc.add_paragraph(f"📌 Sujet : {sujet}")
        doc.add_paragraph(f"📅 Date : {date_intervention.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"🌍 Zone d’intervention : {zone}")
        doc.add_paragraph(f"🏗️ Site : {site}")

        # Sauvegarder en mémoire
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Rapport généré avec succès !")

        # Télécharger le fichier
        st.download_button(
            label="📥 Télécharger le rapport Word",
            data=buffer,
            file_name="rapport.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
